import os
import subprocess
import platform
import logging
import time
import re
from datetime import date
from flask import send_file, flash, redirect, url_for, current_app
from docxtpl import DocxTemplate
from .pdf_utils import convert_docx_to_pdf
from app.utils import turkish_upper

# Loglama yapılandırması
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def post_process_kiralama_docx(docx_path, vergi_no, sozlesme_tarihi_str, kalemler_listesi, makine_kullanim_yeri):
    """
    Üretilen Word dosyasında metin düzeyinde güvenli düzeltmeler yapar.
    """
    try:
        from docx import Document

        docx_doc = Document(docx_path)

        def iter_paragraphs(document):
            for p in document.paragraphs:
                yield p
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            yield p

        for paragraph in iter_paragraphs(docx_doc):
            text = paragraph.text or ""
            if not text:
                continue

            if 'VERGİ KİMLİK NO' in turkish_upper(text) and vergi_no:
                updated = re.sub(
                    r'(VERGİ\s*KİMLİK\s*NO\s*)(X+)',
                    rf'\1{vergi_no}',
                    text,
                    flags=re.IGNORECASE
                )
                if updated != text:
                    paragraph.text = updated
                    text = updated

            if 'Sözleşme Başlangıç Tarihi' in text and sozlesme_tarihi_str:
                updated = re.sub(
                    r'XX\s*\.\s*XX\s*\.\s*20\s*2\s*4',
                    sozlesme_tarihi_str,
                    text
                )
                if updated != text:
                    paragraph.text = updated

            if 'MAKİNE KULLANIM YERİ' in turkish_upper(text) and makine_kullanim_yeri:
                updated = re.sub(
                    r'(MAKİNE\s*KULLANIM\s*YERİ\s*)(.+)$',
                    rf'\1{makine_kullanim_yeri}',
                    text,
                    flags=re.IGNORECASE
                )
                if updated != text:
                    paragraph.text = updated

        docx_doc.save(docx_path)
    except Exception as e:
        logger.warning(f"Word post-process atlandı: {str(e)}")

# Modellerin gerçek konumlarını içe aktarıyoruz
try:
    from app.kiralama.models import Kiralama, KiralamaKalemi
    from app.firmalar.models import Firma
except ImportError as e:
    logger.error(f"Modeller içe aktarılamadı: {e}")

# Blueprint'i aynı klasördeki __init__.py dosyasından alıyoruz
from . import dokumanlar_bp

# Windows ortamında PDF dönüşümü için gerekli
if platform.system() == "Windows":
    import pythoncom

def safe_filename(name):
    """
    Dosya ismindeki tüm riskli karakterleri temizler.
    Harf, rakam, tire (-) ve alt tire (_) DIŞINDAKİ her şeyi '_' yapar.
    Bu yöntem slaş (/) sorununu %100 çözer.
    """
    if not name:
        return "isimsiz_dosya"
    
    # İsmi stringe çevir
    name = str(name).strip()
    
    # Beyaz liste dışındaki her şeyi (boşluk, slaş, nokta vb.) alt tireye çevir
    # re.sub(r'[^a-zA-Z0-9\-_]', ...) -> Alfanumerik, tire ve alt tire haricindekileri eşleştir
    safe_name = re.sub(r'[^a-zA-Z0-9\-_]', '_', name)
    
    # Peş peşe gelen alt tireleri teke indir
    safe_name = re.sub(r'_+', '_', safe_name).strip('_')
    
    return safe_name

def pdf_donustur(docx_path, output_dir):
    """
    Ortak PDF modülünü kullanarak dönüşüm yapar.
    """
    return convert_docx_to_pdf(docx_path, output_dir, logger=logger, timeout_seconds=60)

@dokumanlar_bp.route('/yazdir/form/<int:rental_id>')
def kiralama_formu_yazdir(rental_id):
    """
    Kiralama modellerine tam uyumlu döküman hazırlama rotası.
    """
    try:
        # 1. Veritabanından veriyi çekiyoruz
        kiralama = Kiralama.query.get_or_404(rental_id)
        musteri = kiralama.firma_musteri 
        
        if not musteri:
            flash("Kiralama kaydına bağlı müşteri bulunamadı.", "danger")
            return redirect(url_for('kiralama.index'))

        # --- GENEL SÖZLEŞME KONTROLÜ ---
        gs_no = getattr(musteri, 'sozlesme_no', None)
        gs_trh = getattr(musteri, 'sozlesme_tarihi', None)
        gs_trh_fmt = gs_trh.strftime('%d.%m.%Y') if gs_trh else "BELİRTİLMEDİ"
        if gs_no is None or str(gs_no).strip() == "":
            flash(f"Müşteri ({musteri.firma_adi}) için tanımlı bir Genel Sözleşme bulunamadı.", "danger")
            return redirect(url_for('kiralama.index'))

        # 2. Kalemler Hazırlığı (Döngü)
        kalemler_listesi = []
        genel_toplam = 0
        
        for kalem in kiralama.kalemler:
            gun = (kalem.kiralama_bitis - kalem.kiralama_baslangici).days + 1
            birim_fiyat = float(kalem.kiralama_brm_fiyat or 0)
            nakliye = float(kalem.nakliye_satis_fiyat or 0)
            satir_toplam = (gun * birim_fiyat) + nakliye
            genel_toplam += satir_toplam
            
            if kalem.is_dis_tedarik_ekipman:
                ekipman_adi = kalem.harici_ekipman_marka or "-"
                model = kalem.harici_ekipman_model or "-"
                seri_no = kalem.harici_ekipman_seri_no or "-"
            else:
                if kalem.ekipman:
                    ekipman_adi = kalem.ekipman.marka or "-"
                    model = kalem.ekipman.model or kalem.ekipman.tipi or "-"
                    seri_no = kalem.ekipman.seri_no or "-"
                else:
                    ekipman_adi = "Tanımsız Makine"
                    model = "-"
                    seri_no = "-"

            kalemler_listesi.append({
                'ekipman': ekipman_adi,
                'model': model,
                'seri_no': seri_no,
                'bas_tarih': kalem.kiralama_baslangici.strftime('%d.%m.%Y'),
                'bit_tarih': kalem.kiralama_bitis.strftime('%d.%m.%Y'),
                'gun_sayisi': gun,
                'hizmet_turu': 'MONTAJ',
                'sure': f"{gun} GÜN",
                'birim_fiyat': f"{birim_fiyat:,.2f} TL",
                'tutar': f"{satir_toplam:,.2f} TL",
                'nakliye': f"{nakliye:,.2f} TL",
                'satir_toplam': f"{satir_toplam:,.2f} TL"
            })

        # 3. Dosya Yolları
        base_dir = current_app.root_path 
        template_path = os.path.join(base_dir, 'static', 'templates', 'Kiralama_Formu_TASLAK.docx')
        
        bulut_adi = musteri.bulut_klasor_adi or "Genel_Arsiv"
        
        # Ana çıktı klasörünü oluştur ve normalize et
        output_dir = os.path.abspath(os.path.join(base_dir, 'static', 'arsiv', bulut_adi, 'Formlar'))
        os.makedirs(output_dir, exist_ok=True)
        
        # 4. Dosya Adı Güvenliği (KESİN ÇÖZÜM)
        # PF-2026/0003 -> PF_2026_0003 olacak
        safe_form_no = safe_filename(kiralama.kiralama_form_no)
        docx_filename = f"{safe_form_no}_Form.docx"
        
        # Tam dosya yolunu oluştur ve işletim sistemi formatına zorla
        docx_path = os.path.join(output_dir, docx_filename)

        if not os.path.exists(template_path):
            logger.error(f"Şablon dosyası bulunamadı: {template_path}")
            return f"Taslak bulunamadı: {template_path}", 404

        # 5. Word Şablonunu Doldur
        doc = DocxTemplate(template_path)
        form_tarihi = kiralama.kiralama_olusturma_tarihi.strftime('%d.%m.%Y') if kiralama.kiralama_olusturma_tarihi else date.today().strftime('%d.%m.%Y')
        context = {
            'form_no': kiralama.kiralama_form_no, # Word belgesi içinde / görünebilir, sorun yok
            'gunun_tarihi': form_tarihi,
            'genel_sozlesme_no': gs_no,
            'genel_sozlesme_trh': gs_trh_fmt,
            'makine_kullanim_yeri': kiralama.makine_calisma_adresi or "",
            'musteri_unvan': turkish_upper(musteri.firma_adi),
            'musteri_vergi': f"{musteri.vergi_dairesi or ''} / {musteri.vergi_no or ''}",
            'musteri_vergi_no': musteri.vergi_no or "",
            'musteri_vergi_dairesi': musteri.vergi_dairesi or "",
            'musteri_adres': musteri.iletisim_bilgileri or "",
            'musteri_tel': musteri.telefon or "",
            'kalemler': kalemler_listesi,
            'genel_toplam': f"{genel_toplam:,.2f} TL"
        }
        
        doc.render(context)
        
        # Kaydetme denemesi
        try:
            doc.save(docx_path)
            post_process_kiralama_docx(
                docx_path,
                musteri.vergi_no or "",
                gs_trh_fmt,
                kalemler_listesi,
                kiralama.makine_calisma_adresi or ""
            )
            logger.info(f"Dosya başarıyla kaydedildi: {docx_path}")
        except Exception as e:
            logger.error(f"Kayıt Hatası: {docx_path} - {str(e)}")
            raise e

        # 6. PDF Dönüştürme
        pdf_file = pdf_donustur(docx_path, output_dir)
        
        if pdf_file and os.path.exists(pdf_file):
            try: os.remove(docx_path)
            except: pass
            return send_file(pdf_file, mimetype='application/pdf')
        
        flash("PDF dönüşümü başarısız olduğu için DOCX gönderildi. Sunucuda docx2pdf veya soffice kurulu değil olabilir.", "warning")
        return send_file(docx_path, as_attachment=False)

    except Exception as e:
        logger.error(f"Kiralama Yazdırma Hatası: {str(e)}")
        flash(f"Döküman hazırlanırken hata: {str(e)}", "warning")
        return redirect(url_for('kiralama.index'))